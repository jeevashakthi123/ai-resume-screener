"""
utils/extractor.py
Extracts plain text from PDF and DOCX resume files.
"""

import os
import re
import logging

logger = logging.getLogger(__name__)


def extract_text(filepath: str) -> str:
    """
    Extract text from a PDF or DOCX file.

    Parameters
    ----------
    filepath : str
        Absolute or relative path to the resume file.

    Returns
    -------
    str
        Extracted plain text, or empty string on failure.
    """
    if not os.path.isfile(filepath):
        logger.error("File not found: %s", filepath)
        return ""

    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return _extract_docx(filepath)
    else:
        logger.warning("Unsupported file type: %s", ext)
        return ""


# ── PDF ───────────────────────────────────────────────────────────────────────
def _extract_pdf(filepath: str) -> str:
    """Extract text from PDF using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(filepath)
        return _clean_extracted(text)
    except ImportError:
        logger.warning("pdfminer.six not installed; trying PyPDF2 fallback.")
        return _extract_pdf_pypdf2(filepath)
    except Exception as exc:
        logger.error("PDF extraction failed (%s): %s", filepath, exc)
        return _extract_pdf_pypdf2(filepath)


def _extract_pdf_pypdf2(filepath: str) -> str:
    """Fallback PDF extractor using PyPDF2."""
    try:
        import PyPDF2
        text_parts = []
        with open(filepath, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
        return _clean_extracted("\n".join(text_parts))
    except Exception as exc:
        logger.error("PyPDF2 extraction failed (%s): %s", filepath, exc)
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────
def _extract_docx(filepath: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return _clean_extracted("\n".join(paragraphs))
    except Exception as exc:
        logger.error("DOCX extraction failed (%s): %s", filepath, exc)
        return ""


# ── Cleanup ───────────────────────────────────────────────────────────────────
def _clean_extracted(text: str) -> str:
    """Normalize whitespace and remove non-printable characters."""
    if not text:
        return ""
    # Remove non-printable characters (keep newlines/tabs)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00C0-\u024F]", " ", text)
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)
    return text.strip()
